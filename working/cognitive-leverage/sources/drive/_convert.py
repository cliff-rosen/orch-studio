"""Convert .docx files to .md in place. Preserves paragraphs, headings, bold/italic, and tables."""
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

def inline_runs(paragraph):
    parts = []
    for run in paragraph.runs:
        txt = run.text
        if not txt:
            continue
        if run.bold and run.italic:
            parts.append(f"***{txt}***")
        elif run.bold:
            parts.append(f"**{txt}**")
        elif run.italic:
            parts.append(f"*{txt}*")
        else:
            parts.append(txt)
    return "".join(parts)

def para_to_md(paragraph):
    text = inline_runs(paragraph)
    if not text.strip():
        return ""
    style = paragraph.style.name.lower() if paragraph.style else ""
    if style.startswith("heading 1"):
        return f"# {text}"
    if style.startswith("heading 2"):
        return f"## {text}"
    if style.startswith("heading 3"):
        return f"### {text}"
    if style.startswith("heading 4"):
        return f"#### {text}"
    if style.startswith("heading"):
        return f"##### {text}"
    if style.startswith("title"):
        return f"# {text}"
    if style == "list paragraph" or "list" in style:
        return f"- {text}"
    return text

def table_to_md(table):
    rows = []
    for row in table.rows:
        cells = [" ".join(p.text for p in cell.paragraphs).strip().replace("|", "\\|") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) == 0:
        return ""
    if len(rows) >= 2:
        cols = rows[0].count("|") - 1
        sep = "| " + " | ".join(["---"] * cols) + " |"
        rows.insert(1, sep)
    return "\n".join(rows)

def convert(path_in: Path, path_out: Path):
    doc = Document(str(path_in))
    # Iterate document body in order so tables interleave correctly
    body = doc.element.body
    out = []
    for child in body.iterchildren():
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            # Find the matching paragraph in doc.paragraphs by element identity
            for p in doc.paragraphs:
                if p._element is child:
                    md = para_to_md(p)
                    if md:
                        out.append(md)
                    break
        elif tag == "tbl":
            for t in doc.tables:
                if t._element is child:
                    md = table_to_md(t)
                    if md:
                        out.append(md)
                    break
    text = "\n\n".join(out) + "\n"
    path_out.write_text(text, encoding="utf-8")
    return len(text)

def main():
    root = Path(r"C:/code/orch-studio/working/cognitive-leverage/sources/drive")
    for docx in sorted(root.rglob("*.docx")):
        md = docx.with_suffix(".md")
        size = convert(docx, md)
        print(f"{docx.relative_to(root)} -> {md.relative_to(root)} ({size} bytes)")

if __name__ == "__main__":
    main()
